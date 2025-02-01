import os
import shutil
import time
from pathlib import Path

import cv2
import matplotlib.pyplot as plt
import numpy as np
import uvicorn
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.responses import FileResponse

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.add_middleware(TrustedHostMiddleware, allowed_hosts=["*"])

UPLOAD_DIR = "uploads"
REPORT_DIR = "reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)
def calculate_volume(frame, left, right, roi):
    x, y, w, h = roi
    roi_frame = frame[y:y+h, x:x+w]
    masked_frame = roi_frame[:, left:right] if left and right else roi_frame
    depth_map = estimate_depth(masked_frame)
    material_mask = cv2.threshold(depth_map, 50, 255, cv2.THRESH_BINARY)[1]
    volume = np.sum(material_mask / 255)  # Sum of non-zero pixels
    pixel_to_cubic_cm = 0.5  # Adjust based on calibration
    estimated_volume = volume * pixel_to_cubic_cm
    return estimated_volume

def process_video_volume(video_path, roi):
    cap = cv2.VideoCapture(video_path)
    total_volume = 0
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        frame = cv2.resize(frame, (512, 512))
        
        # Draw ROI bounding box
        x, y, w, h = roi
        cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 2)  # Green box
        
        left, right = detect_wagon_boundaries(frame)
        frame_volume = calculate_volume(frame, left, right, roi)
        total_volume += frame_volume
        
    cap.release()
    return total_volume
def detect_wagon_boundaries(frame):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    
    # Detect vertical lines (theta = 0)
    vertical_lines = cv2.HoughLinesP(edges, 1, np.pi/2, 100, minLineLength=100, maxLineGap=10)
    left, right = None, None
    
    if vertical_lines is not None:
        for line in vertical_lines:
            x1, y1, x2, y2 = line[0]
            if abs(x2 - x1) < 1:  # Vertical line check
                if left is None or x1 < left:
                    left = x1
                if right is None or x1 > right:
                    right = x1
                    
    return left, right

def detect_horizontal_lines(frame):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    
    lines = cv2.HoughLinesP(edges, 1, np.pi / 180, 100, minLineLength=200, maxLineGap=10)
    line_count = 0
    
    if lines is not None:
        for line in lines:
            x1, y1, x2, y2 = line[0]
            if abs(y2 - y1) < 1:
                cv2.line(frame, (x1, y1), (x2, y2), (0, 255, 0), 2)
                line_count += 1
    
    return line_count

def detect_cracks(frame, left_bound, right_bound):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    blurred = cv2.GaussianBlur(enhanced, (5,5), 0)
    edges = cv2.Canny(blurred, 30, 100)
    
    # Morphological closing to connect crack edges
    kernel = np.ones((3,3), np.uint8)
    closed = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel)
    
    contours, _ = cv2.findContours(closed, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    crack_data = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        # Check if contour is within wagon boundaries
        if left_bound is not None and right_bound is not None:
            if x < left_bound or x + w > right_bound:
                continue
                
        area = cv2.contourArea(contour)
        if 50 < area < 5000:
            crack_data.append({
                "area": area,
                "bbox": (x, y, w, h)
            })
            cv2.rectangle(frame, (x, y), (x+w, y+h), (0, 0, 255), 2)
            
    return frame, crack_data
def generate_report(filename, wagon_count, crack_data_list):
    doc = Document()
    doc.add_heading('Advanced Rail Wagon Inspection Report', level=1)
    
    doc.add_paragraph(f'Total Wagons Counted: {wagon_count}')
    doc.add_paragraph(f'Total Cracks Detected: {sum(len(data) for data in crack_data_list)}')
    
    for idx, crack_data in enumerate(crack_data_list, 1):
        doc.add_heading(f'Wagon {idx} Details', level=2)
        
        # Save and insert image
        img_path = os.path.join("report_data", f'wagon_{idx}.png')
        cv2.imwrite(img_path, crack_data['image'])
        doc.add_picture(img_path, width=Inches(4))
        
        # Crack statistics
        doc.add_paragraph(f'Cracks Detected: {len(crack_data["cracks"])}')
        doc.add_paragraph('Crack Sizes (pixels):')
        for i, crack in enumerate(crack_data["cracks"], 1):
            doc.add_paragraph(f'Crack {i}: Area = {crack["area"]:.2f}, Bounding Box = {crack["bbox"]}')
        
        # Add severity visualization
        sizes = [c["area"] for c in crack_data["cracks"]]
        plt.figure()
        plt.hist(sizes, bins=10, color='skyblue')
        plt.title(f'Wagon {idx} Crack Size Distribution')
        plt.xlabel('Crack Area (pixels)')
        plt.ylabel('Frequency')
        chart_path = os.path.join("report_data", f'chart_{idx}.png')
        plt.savefig(chart_path)
        plt.close()
        doc.add_picture(chart_path, width=Inches(4))
    
    report_path = os.path.join(REPORT_DIR, filename)
    doc.save(report_path)
    return report_path
@app.post("/detect")
async def detect(video: UploadFile = File(...)):
    video_path = os.path.join(UPLOAD_DIR, video.filename)
    with open(video_path, "wb") as buffer:
        shutil.copyfileobj(video.file, buffer)
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        print("Error: Could not open video.")
        return
    
    xyzw_list = [(102, 0, 339, 495)]
    x, y, w, h = xyzw_list[0]
    
    wagon_count = 0
    crack_data_list = []
    last_count_time = time.time()
    cooldown = 3
    
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        
        frame = cv2.resize(frame, (512, 512))
        cropped_frame = frame[y:y+h, x:x+w]
        
        # Detect wagon boundaries
        left, right = detect_wagon_boundaries(cropped_frame)
        
        # Process frame
        line_count = detect_horizontal_lines(cropped_frame.copy())
        processed_frame, crack_data = detect_cracks(cropped_frame.copy(), left, right)
        
        if line_count >= 2 and (time.time() - last_count_time) > cooldown:
            wagon_count += 1
            crack_data_list.append({
                "image": processed_frame,
                "cracks": crack_data
            })
            last_count_time = time.time()
            print(f"Wagon {wagon_count} detected with {len(crack_data)} cracks")
        
        # Draw boundaries
        if left and right:
            cv2.line(processed_frame, (left, 0), (left, h), (255,0,0), 2)
            cv2.line(processed_frame, (right, 0), (right, h), (255,0,0), 2)
        
        # cv2.imshow("Processed Frame", processed_frame)
        # if cv2.waitKey(1) & 0xFF == ord('q'):
        #     break
    
    # generate_report(wagon_count, crack_data_list)
    cap.release()
    report_filename = f"report_{video.filename}.docx"
    report_path = generate_report(report_filename, wagon_count, crack_data_list)
    return {"report_url": f"/download/{report_filename}"}

@app.get("/download/{report_name}")
def download_report(report_name: str):
    report_path = Path(REPORT_DIR) / report_name
    if not report_path.exists():
        return {"error": "Report not found"}
    
    return FileResponse(report_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=report_name)

@app.post("/getvolume")
async def process_video_api(video: UploadFile = File(...)):
    video_path = f"uploads/{video.filename}"
    
    # Save the uploaded video
    with open(video_path, "wb") as buffer:
        shutil.copyfileobj(video.file, buffer)
    
    roi_coords = (102, 0, 339, 495)  # Example ROI values
    total_volume = process_video_volume(video_path, roi_coords)  # Call existing function
    
    return {"message": "Processing complete", "estimated_volume": total_volume}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
